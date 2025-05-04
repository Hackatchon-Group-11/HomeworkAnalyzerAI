from vertexai.language_models import TextGenerationModel
from config import Config
from docx import Document
from markdown2 import markdown
from weasyprint import HTML

class QuestionAnswerAnalyzer:
    def __init__(self):
        self.model = TextGenerationModel.from_pretrained(Config.MODEL_NAME)

    def analyze_qa_pairs(self, text: str) -> list:
        prompt = f"""
Aşağıdaki metinde soru-cevap çiftleri bulunmaktadır. Her bir çift için:
- Soruyu ve cevabı tanımla
- Cevabın uygun olup olmadığını değerlendir (doğru, eksik, yanlış)
- Eksik veya hatalıysa doğru cevabı öner

Metin:
{text}

Çıktı formatı:
[
  {{
    "soru": "...",
    "verilen_cevap": "...",
    "degerlendirme": "doğru / eksik / yanlış",
    "duzeltme": "..." (eğer gerekliyse)
  }},
  ...
]
"""

        try:
            response = self.model.predict(prompt, temperature=0.3, max_output_tokens=2048)
            return self._parse_response(response.text)
        except Exception as e:
            print(f"[QA Analysis Error] {e}")
            return []

    def _parse_response(self, response: str) -> list:
        import json
        try:
            # Bazı durumlarda model düzgün formatta dönmeyebilir. JSON olarak parse etmeye çalışıyoruz.
            return json.loads(response)
        except json.JSONDecodeError:
            # Alternatif olarak satır bazlı parse edilebilir
            print("[Warning] JSON parse failed, raw response:\n", response)
            return []

    def to_markdown(self, qa_list: list) -> str:
        md_output = "# 📘 Soru-Cevap Analizi\n\n"
        for i, item in enumerate(qa_list, 1):
            md_output += f"## {i}. Soru\n"
            md_output += f"- **❓ Soru:** {item.get('soru', '').strip()}\n"
            md_output += f"- **✅ Verilen Cevap:** {item.get('verilen_cevap', '').strip()}\n"
            md_output += f"- **🧠 Değerlendirme:** `{item.get('degerlendirme', '').strip()}`\n"
            if item.get("duzeltme"):
                md_output += f"- **💡 Önerilen Cevap:** {item['duzeltme'].strip()}\n"
            md_output += "\n"
        return md_output

    def to_docx(self, qa_list: list, output_path: str):
        doc = Document()
        doc.add_heading("Soru-Cevap Analizi", level=1)
        for i, item in enumerate(qa_list, 1):
            doc.add_heading(f"{i}. Soru", level=2)
            doc.add_paragraph(f" Soru: {item.get('soru', '').strip()}")
            doc.add_paragraph(f" Verilen Cevap: {item.get('verilen_cevap', '').strip()}")
            doc.add_paragraph(f" Değerlendirme: {item.get('degerlendirme', '').strip()}")
            if item.get("duzeltme"):
                doc.add_paragraph(f"💡 Önerilen Cevap: {item['duzeltme'].strip()}")
        doc.save(output_path)

    def to_pdf(self, qa_list: list, output_path: str):
        html_content = markdown(self.to_markdown(qa_list))
        HTML(string=html_content).write_pdf(output_path)